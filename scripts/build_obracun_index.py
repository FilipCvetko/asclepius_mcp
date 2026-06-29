#!/usr/bin/env python3
"""Build the isolated `obracun_rules` ChromaDB collection (outpatient billing rules, v1).

Sources (current versions, from data/billing_sources.json):
  • Uredba o programih storitev OZZ 2026   — OCR text (scanned PDF) via <unid>_ocr.json
  • Navodilo o beleženju/obračunu          — latest 'Vsebinsko navodilo v## čistopis' (docx in zip)
  • Standardi kodiranja                    — consolidated docx
  • Billing Q&A                            — docx (skip scanned pdf duplicates)

Each source is extracted to page-level text, chunked (reuse sliding_window_chunk), and
embedded into a collection isolated from the broad e-gradiva index. Run locally.
"""
import io
import json
import re
import sys
import zipfile
from pathlib import Path

import fitz  # PyMuPDF

sys.path.insert(0, str(Path(__file__).parent.parent))
from build_egradiva_index import sliding_window_chunk  # noqa: E402

DATA = Path(__file__).parent.parent / "data"
SOURCES = DATA / "billing_sources.json"
CHROMA_DIR = Path(__import__("os").environ.get("CHROMADB_PATH", DATA / "chromadb"))
COLLECTION_NAME = "obracun_rules"
EMBEDDING_MODEL = "paraphrase-multilingual-MiniLM-L12-v2"

TITLES = {
    "uredba_programi": "Uredba o programih storitev OZZ (2026)",
    "navodilo_obracun": "Navodilo o beleženju in obračunavanju – Vsebinsko navodilo (čistopis)",
    "standardi_kodiranja": "Standardi kodiranja",
    "billing_qa": "Navodilo za obračun – vprašanja in odgovori",
}


def _docx_text(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:                       # tables carry codes/points — keep them
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _pdf_pages(path: str):
    d = fitz.open(path)
    return [(i + 1, d[i].get_text("text").strip()) for i in range(len(d))
            if len(d[i].get_text("text").strip()) > 40]


def _latest_navodilo_docx(files: list) -> bytes | None:
    """Across the Navodilo zips, return the highest-version 'čistopis' Vsebinsko navodilo docx."""
    best = (-1, None)
    for f in files:
        if f["ext"] != "zip":
            continue
        try:
            zf = zipfile.ZipFile(f["path"])
        except Exception:
            continue
        for name in zf.namelist():
            if not name.lower().endswith(".docx"):
                continue
            m = re.search(r"vsebinsko\s*navodilo\s*v(\d+)", name, re.I)
            if not m:
                continue
            ver = int(m.group(1))
            consolidated = bool(re.search(r"istopis", name, re.I))  # čistopis/cistopis
            score = ver * 10 + (1 if consolidated else 0)
            if score > best[0]:
                best = (score, (zf, name))
    if best[1]:
        zf, name = best[1]
        print(f"   navodilo: using '{name}' (v-score {best[0]})")
        return zf.read(name)
    return None


def pages_for_source(s: dict):
    """Return list of (page_number, text) for one source, picking current/clean files."""
    role = s["role"]
    out = []

    if role == "uredba_programi":
        ocr = DATA / "billing_files" / f"{s['unid']}_ocr.json"
        if ocr.exists():
            for p in json.loads(ocr.read_text(encoding="utf-8")):
                if len(p["text"]) > 40:
                    out.append((p["page"], p["text"]))
        else:
            print("   WARNING: Uredba OCR json missing — run scripts/ocr_uredba.py first")
        return out

    if role == "navodilo_obracun":
        # Only the latest 'čistopis' Vsebinsko navodilo from the zips — the loose .doc
        # attachments are older versions/annexes (avoid version sprawl).
        data = _latest_navodilo_docx(s["files"])
        if data:
            out.append((1, _docx_text(data)))
        else:
            print("   WARNING: no čistopis Vsebinsko navodilo found in zips")
        return out

    # standardi_kodiranja / billing_qa: prefer docx; else text-bearing pdf
    docx_files = [f for f in s["files"] if f["ext"] in ("docx", "doc")]
    used = False
    for f in (docx_files[:1] if role == "standardi_kodiranja" else docx_files):
        try:
            out.append((1, _docx_text(Path(f["path"]).read_bytes())))
            used = True
        except Exception:
            pass
    if not used:
        for f in s["files"]:
            if f["ext"] == "pdf":
                pg = _pdf_pages(f["path"])
                if pg:
                    out.extend(pg)
                    break
    return out


def main() -> None:
    import chromadb
    from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction

    sources = json.loads(SOURCES.read_text(encoding="utf-8"))
    embed = SentenceTransformerEmbeddingFunction(model_name=EMBEDDING_MODEL)
    client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    try:
        client.delete_collection(COLLECTION_NAME)
    except Exception:
        pass
    coll = client.create_collection(name=COLLECTION_NAME, embedding_function=embed,
                                    metadata={"hnsw:space": "cosine"})

    ids, docs, metas = [], [], []
    for s in sources:
        title = TITLES.get(s["role"], s["title"])
        print(f"\n• {s['role']}  ({s.get('date','')})  → {title}")
        n0 = len(ids)
        for page, text in pages_for_source(s):
            for ch in sliding_window_chunk(text, title, s["role"], s["unid"],
                                           s.get("files", [{}])[0].get("url", ""), page):
                ids.append(f"obracun_{s['unid']}_{len(ids)}")
                docs.append(ch["text"])
                metas.append({
                    "doc_title": title, "role": s["role"], "domain": s.get("domain", "outpatient"),
                    "date": s.get("date", ""), "page": page,
                    "source_url": ch["metadata"].get("file_url", ""),
                    "source": "ZZZS obračun",
                })
        print(f"   chunks: {len(ids) - n0}")

    print(f"\nEmbedding {len(ids)} chunks...")
    for i in range(0, len(ids), 100):
        coll.add(ids=ids[i:i+100], documents=docs[i:i+100], metadatas=metas[i:i+100])
    print(f"Done: collection '{COLLECTION_NAME}' = {coll.count()} chunks at {CHROMA_DIR}")


if __name__ == "__main__":
    main()
